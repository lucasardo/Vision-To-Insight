import streamlit as st
import base64
import requests
from langchain_openai import AzureChatOpenAI
import json
from dotenv import load_dotenv
import os
from io import BytesIO
from PIL import Image
from docx import Document
from docx.shared import Inches
load_dotenv(os.path.join("credentials.env"))

#### SETUP

url = os.environ['FUNCTION_APP_URL']
headers = {"Content-Type": "application/json"}

azure_endpoint = os.environ['GLOBAL_AZURE_ENDPOINT']
openai_api_key = os.environ['GLOBAL_OPENAI_API_KEY']
openai_deployment_name = os.environ['GLOBAL_GPT_DEPLOYMENT_NAME']
openai_api_version = os.environ['GLOBAL_OPENAI_API_VERSION']

llm = AzureChatOpenAI(
    azure_endpoint=azure_endpoint,
    api_key=openai_api_key,
    deployment_name=openai_deployment_name,
    api_version=openai_api_version
)

#### APP

st.markdown("<h1 style='text-align: center; color: white;'>VISION TO INSIGHT - WSP Italy</h1>", unsafe_allow_html=True)
st.write("Un modello di AI multimodale per analizzare immagini di cantiere e generare automaticamente report tecnici")

# Initialize session state variables
if "uploaded_file" not in st.session_state:
    st.session_state.uploaded_file = None

if "analysis_content" not in st.session_state:
    st.session_state.analysis_content = None

if "json_data" not in st.session_state:
    st.session_state.json_data = None

# File uploader
uploaded_file = st.file_uploader("Scegli un'immagine...", type="jpg")

if uploaded_file:
    st.session_state.uploaded_file = uploaded_file
    st.session_state.analysis_content = None  # Reset analysis when a new image is uploaded
    st.session_state.json_data = None  # Reset JSON data as well

# Perform analysis only if an image is uploaded and no analysis exists in the state
if st.session_state.uploaded_file and st.session_state.analysis_content is None:
    with st.spinner("Analisi dell'immagine in corso..."):
        try:
            # Prepare the image data for the API
            image_data = base64.b64encode(st.session_state.uploaded_file.read()).decode()
            payload = {"image_data": image_data}
            print(payload)
            # Make the API request
            response = requests.post(url, data=json.dumps(payload), headers=headers)
            response.raise_for_status()

            # Parse the JSON response
            st.session_state.json_data = response.json()

            # Generate analysis content using the LLM
            analysis_prompt = f"""
            Sei un ingegnere civile e ti è stato affidato il compito di scrivere una sezione di una relazione tecnica sull'ispezione dei pozzetti.
            Il nome della sezione è “Analisi Descrittiva”.
            In questa sezione, devi descrivere accuratamente ogni pozzetto in base ai dati di ispezione forniti.
            Il risultato deve essere un paragrafo di testo per ogni pozzetto, che descriva i dati di ispezione. Non utilizzare elenchi puntati.
            Dopo la descrizione testuale, includi una tabella che riassume i dati di ispezione.
            Questo è lo schema JSON che contiene i dati dell'ispezione dei tombini:

            {st.session_state.json_data}
            
            Esempio di output:
            
            ## Analisi Descrittiva
            Il pozzetto ispezionato appartiene al sistema idrico, in quanto è utilizzato per la gestione dell'acqua. Durante l'ispezione, è stato rilevato che il pozzetto non è associato ad alcun sistema di comunicazione, elettricità, termico, gas, acque reflue, né a sistemi di acque meteoriche, ma esclusivamente a quello idrico. Il pozzetto in questione è identificato come 'wsp_manhole' e presenta una forma rettangolare, con dimensioni approssimative di 60 centimetri per 60 centimetri. Le dimensioni e la forma sono state verificate durante l'ispezione e risultano conformi agli standard previsti per questo tipo di infrastruttura. Non sono stati rilevati danni visibili o anomalie che potrebbero compromettere la funzionalità del pozzetto.

            Caratteristica	Dati
            Utility	Acqua
            Sistema di Comunicazione	Non presente
            Sistema Elettrico	Non presente
            Sistema Termico	Non presente
            Sistema Gas	Non presente
            Sistema Acque Reflue	Non presente
            Sistema Acque Meteoriche	Non presente
            Sistema Idrico	wsp_manhole
            Forma	Rettangolare
            Dimensioni	Circa 60cm x 60cm
        
            """
            analysis = llm.invoke(analysis_prompt)
            st.session_state.analysis_content = analysis.content

        except requests.exceptions.RequestException as e:
            st.error(f"Errore durante l'analisi: {e}")
            st.stop()

# Display analysis content
if st.session_state.analysis_content:
    st.write(st.session_state.analysis_content)

# Provide report download option
if st.session_state.analysis_content:
    document = Document()
    document.add_heading("Analisi Descrittiva", level=1)

    # Add analysis content to the Word document
    for paragraph in st.session_state.analysis_content.split("\n\n"):
        document.add_paragraph(paragraph)

    # Add the uploaded image to the document
    image = Image.open(st.session_state.uploaded_file)
    image_path = BytesIO()
    image.save(image_path, format="PNG")
    image_path.seek(0)
    document.add_heading("Immagine dell'ispezione", level=2)
    document.add_picture(image_path, width=Inches(4))

    # Save the document to a buffer
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    st.download_button(
        label="📄",  # Unicode for a page or document icon
        data=buffer,
        file_name="V2I_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        help="Scarica il report in formato Word"
    )

# Display image preview and JSON data
if st.session_state.uploaded_file:
    col1, col2 = st.columns(2)
    with col1:
        st.image(st.session_state.uploaded_file, caption="Anteprima immagine", use_column_width=True)
    with col2:
        if st.session_state.json_data:
            with st.expander("Mostra lo schema JSON"):
                st.json(st.session_state.json_data)